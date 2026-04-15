"""
Spine Engine — Word Parser.

Parses Word documents (.docx) — typically BRDs, compliance manuals,
underwriting guides. Uses python-docx for real parsing.
"""

from __future__ import annotations

import io
import logging

from nexus_sdk.events import fire_stub_alert

from ..models import ExtractedTable

logger = logging.getLogger("nexus.spine.parsers.word")

# Module-level event bus and stub tracking
_event_bus = None
_stub_counts: dict[str, int] = {}


def set_event_bus(bus) -> None:
    """Called by SpineEngine.on_startup() to inject the event bus reference."""
    global _event_bus
    _event_bus = bus


class WordParser:
    """Parse Word documents (BRDs, compliance manuals)."""

    @staticmethod
    def parse(content: bytes, document_id: str) -> dict:
        try:
            import docx  # type: ignore[import-not-found]
            return WordParser._parse_with_docx(content, document_id)
        except ImportError:
            return WordParser._stub_parse(content, document_id)

    @staticmethod
    def _parse_with_docx(content: bytes, document_id: str) -> dict:
        import docx  # type: ignore[import-not-found]

        doc = docx.Document(io.BytesIO(content))
        paragraphs = []
        full_text = ""
        current_section = "Introduction"
        tables = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Detect headings
            if para.style and para.style.name and "Heading" in para.style.name:
                current_section = text

            paragraphs.append({
                "text": text,
                "style": para.style.name if para.style else "Normal",
                "section": current_section,
            })
            full_text += text + "\n"

        # Extract tables from Word doc
        for idx, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)

            if rows:
                headers = rows[0]
                tables.append(ExtractedTable(
                    document_id=document_id,
                    headers=headers,
                    rows=rows[1:],
                    row_count=len(rows) - 1,
                    col_count=len(headers),
                    metadata={"table_index": idx},
                ))

        return {
            "paragraphs": paragraphs,
            "tables": tables,
            "full_text": full_text,
            "page_count": max(1, len(full_text) // 3000),
        }

    @staticmethod
    def _stub_parse(content: bytes, document_id: str) -> dict:
        _stub_counts["word"] = _stub_counts.get("word", 0) + 1
        logger.warning("spine: Word parser stub fallback #%d", _stub_counts["word"])
        fire_stub_alert(
            _event_bus, "spine", "word_parser",
            fallback_count=_stub_counts["word"],
            reason="python-docx not installed",
        )
        return {
            "paragraphs": [],
            "tables": [],
            "full_text": "[Stub Word doc — install python-docx for real parsing]",
            "page_count": 1,
        }
